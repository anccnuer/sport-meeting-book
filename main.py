import os
import sys

# 检查是否有命令行参数，如果有则使用命令行模式
if len(sys.argv) > 1 and sys.argv[1] == '--cli':
    # 命令行模式
    import sqlite3
    import pandas as pd
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    class SportsMeetManager:
        def __init__(self, excel_path, db_path='sports_meet.db'):
            self.excel_path = excel_path
            self.db_path = db_path
            self.grade_sheets = {}
            self.events_by_grade = {}
            self.initialize_database()
        
        def initialize_database(self):
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            cursor.execute('''
            CREATE TABLE IF NOT EXISTS students (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                grade INTEGER,
                class_name TEXT,
                gender TEXT,
                number TEXT,
                name TEXT
            )
            ''')
            cursor.execute('''
            CREATE TABLE IF NOT EXISTS events (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT UNIQUE
            )
            ''')
            cursor.execute('''
            CREATE TABLE IF NOT EXISTS student_events (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                student_id INTEGER,
                event_id INTEGER,
                FOREIGN KEY (student_id) REFERENCES students (id),
                FOREIGN KEY (event_id) REFERENCES events (id)
            )
            ''')
            conn.commit()
            conn.close()
        
        def parse_excel(self):
            xls = pd.ExcelFile(self.excel_path)
            for sheet_name in xls.sheet_names:
                if '年级' in sheet_name:
                    grade = self.extract_grade(sheet_name)
                    if grade:
                        self.grade_sheets[grade] = sheet_name
                        df = pd.read_excel(self.excel_path, sheet_name=sheet_name)
                        self.parse_sheet(df, grade)
        
        def extract_grade(self, sheet_name):
            for char in sheet_name:
                if char.isdigit():
                    return int(char)
            return None
        
        def parse_sheet(self, df, grade):
            # 找到表头行
            header_row = 0
            for i, row in df.iterrows():
                if isinstance(row.iloc[0], str) and '号码' in row.iloc[0]:
                    header_row = i
                    break
            
            # 提取事件名称
            event_row = header_row + 1
            events = []
            for col in df.columns[2:]:
                event_name = str(df.iloc[event_row, df.columns.get_loc(col)]).strip()
                if event_name and event_name != 'nan':
                    events.append(event_name)
            
            self.events_by_grade[grade] = events
            
            # 提取班级和性别信息
            class_info = str(df.iloc[1, 0]).strip()
            class_name = ''
            gender = ''
            if '班级：' in class_info:
                class_name = class_info.split('班级：')[1].split('参赛组别：')[0].strip()
            if '参赛组别：' in class_info:
                gender = class_info.split('参赛组别：')[1].strip()
            
            # 提取学生信息
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            for i in range(event_row + 1, len(df)):
                row = df.iloc[i]
                number = str(row.iloc[0]).strip()
                name = str(row.iloc[1]).strip()
                
                if not number or number == 'nan' or not name or name == 'nan':
                    continue
                
                # 插入学生信息
                cursor.execute('''
                INSERT OR IGNORE INTO students (grade, class_name, gender, number, name)
                VALUES (?, ?, ?, ?, ?)
                ''', (grade, class_name, gender, number, name))
                
                # 获取学生ID
                cursor.execute('SELECT id FROM students WHERE number = ?', (number,))
                student_id = cursor.fetchone()[0]
                
                # 处理参赛项目
                for j, event_name in enumerate(events):
                    col_index = 2 + j
                    if col_index < len(df.columns):
                        value = str(row.iloc[col_index]).strip()
                        if value and value != 'nan':
                            # 插入事件
                            cursor.execute('INSERT OR IGNORE INTO events (name) VALUES (?)', (event_name,))
                            # 获取事件ID
                            cursor.execute('SELECT id FROM events WHERE name = ?', (event_name,))
                            event_id = cursor.fetchone()[0]
                            # 插入学生-事件关系
                            cursor.execute('''
                            INSERT OR IGNORE INTO student_events (student_id, event_id)
                            VALUES (?, ?)
                            ''', (student_id, event_id))
            
            conn.commit()
            conn.close()
        
        def generate_order_book(self, output_path='秩序册.docx'):
            doc = Document()
            
            # 设置字体
            for style in doc.styles:
                if style.name in ['Normal', 'Heading 1', 'Heading 2', 'Heading 3']:
                    style.font.name = '微软雅黑'
                    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            
            # 添加标题
            title = doc.add_heading('运动会秩序册', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.runs[0].font.size = Pt(24)
            
            # 按年级生成秩序
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            for grade in sorted(self.grade_sheets.keys()):
                grade_heading = doc.add_heading(f'{grade}年级', level=1)
                
                for event_name in self.events_by_grade[grade]:
                    event_heading = doc.add_heading(event_name, level=2)
                    
                    # 查询参加该项目的学生
                    cursor.execute('''
                    SELECT s.number, s.name, s.class_name, s.gender
                    FROM students s
                    JOIN student_events se ON s.id = se.student_id
                    JOIN events e ON se.event_id = e.id
                    WHERE s.grade = ? AND e.name = ?
                    ORDER BY s.number
                    ''', (grade, event_name))
                    
                    students = cursor.fetchall()
                    if students:
                        table = doc.add_table(rows=1, cols=4)
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = '号码'
                        hdr_cells[1].text = '姓名'
                        hdr_cells[2].text = '班级'
                        hdr_cells[3].text = '性别'
                        
                        for student in students:
                            row_cells = table.add_row().cells
                            row_cells[0].text = student[0]
                            row_cells[1].text = student[1]
                            row_cells[2].text = student[2]
                            row_cells[3].text = student[3]
                    
                    doc.add_paragraph()
            
            conn.close()
            doc.save(output_path)
            print(f'秩序册已生成：{output_path}')

    if __name__ == '__main__':
        excel_path = 'data/origin_data.xlsx'
        if len(sys.argv) > 2:
            excel_path = sys.argv[2]
        manager = SportsMeetManager(excel_path)
        manager.parse_excel()
        manager.generate_order_book()
else:
    # GUI模式
    try:
        from src.gui.main_window import main
        if __name__ == '__main__':
            main()
    except ImportError as e:
        print(f"无法启动GUI: {e}")
        print("尝试使用命令行模式...")
        # 回退到命令行模式
        import sqlite3
        import pandas as pd
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        class SportsMeetManager:
            def __init__(self, excel_path, db_path='sports_meet.db'):
                self.excel_path = excel_path
                self.db_path = db_path
                self.grade_sheets = {}
                self.events_by_grade = {}
                self.initialize_database()
            
            def initialize_database(self):
                conn = sqlite3.connect(self.db_path)
                cursor = conn.cursor()
                cursor.execute('''
                CREATE TABLE IF NOT EXISTS students (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    grade INTEGER,
                    class_name TEXT,
                    gender TEXT,
                    number TEXT,
                    name TEXT
                )
                ''')
                cursor.execute('''
                CREATE TABLE IF NOT EXISTS events (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    name TEXT UNIQUE
                )
                ''')
                cursor.execute('''
                CREATE TABLE IF NOT EXISTS student_events (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    student_id INTEGER,
                    event_id INTEGER,
                    FOREIGN KEY (student_id) REFERENCES students (id),
                    FOREIGN KEY (event_id) REFERENCES events (id)
                )
                ''')
                conn.commit()
                conn.close()
            
            def parse_excel(self):
                xls = pd.ExcelFile(self.excel_path)
                for sheet_name in xls.sheet_names:
                    if '年级' in sheet_name:
                        grade = self.extract_grade(sheet_name)
                        if grade:
                            self.grade_sheets[grade] = sheet_name
                            df = pd.read_excel(self.excel_path, sheet_name=sheet_name)
                            self.parse_sheet(df, grade)
            
            def extract_grade(self, sheet_name):
                for char in sheet_name:
                    if char.isdigit():
                        return int(char)
                return None
            
            def parse_sheet(self, df, grade):
                # 找到表头行
                header_row = 0
                for i, row in df.iterrows():
                    if isinstance(row.iloc[0], str) and '号码' in row.iloc[0]:
                        header_row = i
                        break
                
                # 提取事件名称
                event_row = header_row + 1
                events = []
                for col in df.columns[2:]:
                    event_name = str(df.iloc[event_row, df.columns.get_loc(col)]).strip()
                    if event_name and event_name != 'nan':
                        events.append(event_name)
                
                self.events_by_grade[grade] = events
                
                # 提取班级和性别信息
                class_info = str(df.iloc[1, 0]).strip()
                class_name = ''
                gender = ''
                if '班级：' in class_info:
                    class_name = class_info.split('班级：')[1].split('参赛组别：')[0].strip()
                if '参赛组别：' in class_info:
                    gender = class_info.split('参赛组别：')[1].strip()
                
                # 提取学生信息
                conn = sqlite3.connect(self.db_path)
                cursor = conn.cursor()
                
                for i in range(event_row + 1, len(df)):
                    row = df.iloc[i]
                    number = str(row.iloc[0]).strip()
                    name = str(row.iloc[1]).strip()
                    
                    if not number or number == 'nan' or not name or name == 'nan':
                        continue
                    
                    # 插入学生信息
                    cursor.execute('''
                    INSERT OR IGNORE INTO students (grade, class_name, gender, number, name)
                    VALUES (?, ?, ?, ?, ?)
                    ''', (grade, class_name, gender, number, name))
                    
                    # 获取学生ID
                    cursor.execute('SELECT id FROM students WHERE number = ?', (number,))
                    student_id = cursor.fetchone()[0]
                    
                    # 处理参赛项目
                    for j, event_name in enumerate(events):
                        col_index = 2 + j
                        if col_index < len(df.columns):
                            value = str(row.iloc[col_index]).strip()
                            if value and value != 'nan':
                                # 插入事件
                                cursor.execute('INSERT OR IGNORE INTO events (name) VALUES (?)', (event_name,))
                                # 获取事件ID
                                cursor.execute('SELECT id FROM events WHERE name = ?', (event_name,))
                                event_id = cursor.fetchone()[0]
                                # 插入学生-事件关系
                                cursor.execute('''
                                INSERT OR IGNORE INTO student_events (student_id, event_id)
                                VALUES (?, ?)
                                ''', (student_id, event_id))
                
                conn.commit()
                conn.close()
            
            def generate_order_book(self, output_path='秩序册.docx'):
                doc = Document()
                
                # 设置字体
                for style in doc.styles:
                    if style.name in ['Normal', 'Heading 1', 'Heading 2', 'Heading 3']:
                        style.font.name = '微软雅黑'
                        style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                
                # 添加标题
                title = doc.add_heading('运动会秩序册', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title.runs[0].font.size = Pt(24)
                
                # 按年级生成秩序
                conn = sqlite3.connect(self.db_path)
                cursor = conn.cursor()
                
                for grade in sorted(self.grade_sheets.keys()):
                    grade_heading = doc.add_heading(f'{grade}年级', level=1)
                    
                    for event_name in self.events_by_grade[grade]:
                        event_heading = doc.add_heading(event_name, level=2)
                        
                        # 查询参加该项目的学生
                        cursor.execute('''
                        SELECT s.number, s.name, s.class_name, s.gender
                        FROM students s
                        JOIN student_events se ON s.id = se.student_id
                        JOIN events e ON se.event_id = e.id
                        WHERE s.grade = ? AND e.name = ?
                        ORDER BY s.number
                        ''', (grade, event_name))
                        
                        students = cursor.fetchall()
                        if students:
                            table = doc.add_table(rows=1, cols=4)
                            hdr_cells = table.rows[0].cells
                            hdr_cells[0].text = '号码'
                            hdr_cells[1].text = '姓名'
                            hdr_cells[2].text = '班级'
                            hdr_cells[3].text = '性别'
                            
                            for student in students:
                                row_cells = table.add_row().cells
                                row_cells[0].text = student[0]
                                row_cells[1].text = student[1]
                                row_cells[2].text = student[2]
                                row_cells[3].text = student[3]
                        
                        doc.add_paragraph()
                
                conn.close()
                doc.save(output_path)
                print(f'秩序册已生成：{output_path}')

        if __name__ == '__main__':
            excel_path = 'data/origin_data.xlsx'
            manager = SportsMeetManager(excel_path)
            manager.parse_excel()
            manager.generate_order_book()
