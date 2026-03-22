from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

class WordWriter:
    def __init__(self):
        self.doc = Document()
    
    def set_font(self):
        """设置文档字体"""
        for style in self.doc.styles:
            if style.name in ['Normal', 'Heading 1', 'Heading 2', 'Heading 3']:
                style.font.name = '微软雅黑'
                style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    def add_title(self, title, level=0):
        """添加标题"""
        heading = self.doc.add_heading(title, level)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if level == 0:
            heading.runs[0].font.size = Pt(24)
        return heading
    
    def add_paragraph(self, text=''):
        """添加段落"""
        return self.doc.add_paragraph(text)
    
    def add_table(self, rows, cols):
        """添加表格"""
        return self.doc.add_table(rows=rows, cols=cols)
    
    def save(self, file_path):
        """保存文档"""
        self.doc.save(file_path)
